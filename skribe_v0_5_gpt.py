from openai import OpenAI
import streamlit as st
from docx import Document
import PyPDF2
import io

# Set your OpenAI API key securely
client = OpenAI(api_key="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")  # 🔐 Replace with your real key

# --- Utility Functions ---
def extract_text_from_file(uploaded_file):
    if uploaded_file.name.endswith('.pdf'):
        reader = PyPDF2.PdfReader(uploaded_file)
        text = ''.join(page.extract_text() for page in reader.pages if page.extract_text())
    elif uploaded_file.name.endswith('.docx'):
        doc = Document(uploaded_file)
        text = '\n'.join([para.text for para in doc.paragraphs])
    elif uploaded_file.name.endswith('.txt'):
        text = uploaded_file.read().decode('utf-8')
    else:
        text = ""
    return text

def call_openai_gpt(prompt):
    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error calling OpenAI: {e}"

def generate_cover_letter_gpt(job_desc, resume_text):
    prompt = f"""
Write a personalized cover letter for the following job description and resume.
Keep it professional and tailored for a mid-senior business analyst role.
Limit to 300 words.

Job Description:
{job_desc}

Resume:
{resume_text}
"""
    return call_openai_gpt(prompt)

def tailor_resume_gpt(job_desc, resume_text):
    prompt = f"""
Improve and tailor the resume below to better align with the job description.
Update the summary, key skills, and experience descriptions to reflect the role requirements.

Job Description:
{job_desc}

Resume:
{resume_text}
"""
    return call_openai_gpt(prompt)

def save_as_docx(text):
    doc = Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Streamlit UI ---
st.set_page_config(page_title="Skribe AI – GPT Assistant", layout="centered")
st.title("📄 Skribe v0.5 – GPT-Powered Job Application Assistant")

# Step 1: Job Description
st.header("Step 1: Provide Job Description")
input_method = st.radio("How would you like to provide the job description?", ["Paste Text", "Upload File"])
job_description = ""

if input_method == "Paste Text":
    job_description = st.text_area("Paste job description here:")
else:
    job_file = st.file_uploader("Upload job description file (.pdf, .docx, .txt)", type=["txt", "pdf", "docx"])
    if job_file:
        job_description = extract_text_from_file(job_file)

# Step 2: Resume Upload
st.header("Step 2: Upload Your Resume")
resume_file = st.file_uploader("Upload your resume (.pdf or .docx)", type=["pdf", "docx"])
resume_text = ""
if resume_file:
    resume_text = extract_text_from_file(resume_file)

# Step 3: Generate
if st.button("✨ Generate AI-Personalized Cover Letter & Resume"):
    if not job_description or not resume_text:
        st.error("Please provide both a job description and a resume.")
    else:
        with st.spinner("Generating personalized documents using GPT-3.5 Turbo..."):
            cover_letter = generate_cover_letter_gpt(job_description, resume_text)
            tailored_resume = tailor_resume_gpt(job_description, resume_text)

        st.success("✅ Your documents are ready!")

        st.download_button("📥 Download Cover Letter (DOCX)", save_as_docx(cover_letter), "cover_letter.docx")
        st.download_button("📥 Download Tailored Resume (DOCX)", save_as_docx(tailored_resume), "tailored_resume.docx")